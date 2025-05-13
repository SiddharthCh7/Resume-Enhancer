import os, requests, json, uuid, re

from fastapi import FastAPI, File, UploadFile, Form, Request, Response
from fastapi.responses import HTMLResponse, FileResponse, RedirectResponse
from fastapi.templating import Jinja2Templates

import fitz
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER

from pathlib import Path
from dotenv import load_dotenv

from starlette.middleware.base import BaseHTTPMiddleware
from sqlalchemy import create_engine, Column, String
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker

load_dotenv()
app = FastAPI()

# Database setup
DATABASE_URL = "sqlite:///./sessions.db"
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class SessionModel(Base):
    __tablename__ = "sessions"
    id = Column(String, primary_key=True)
    data = Column(String)

Base.metadata.create_all(bind=engine)

# Custom session middleware
class CustomSessionMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        session_id = request.cookies.get("session_id")
        
        # Create a new session if one doesn't exist
        if not session_id:
            session_id = str(uuid.uuid4())
            request.state.session = {}
            request.state.new_session = True
        else:
            # Load existing session from database
            db = SessionLocal()
            try:
                db_session = db.query(SessionModel).filter(SessionModel.id == session_id).first()
                if db_session:
                    request.state.session = json.loads(db_session.data)
                else:
                    request.state.session = {}
                    request.state.new_session = True
            finally:
                db.close()
        
        # Process the request
        response = await call_next(request)
        
        # Save session to database
        db = SessionLocal()
        try:
            session_data = json.dumps(request.state.session)
            db_session = db.query(SessionModel).filter(SessionModel.id == session_id).first()
            if db_session:
                db_session.data = session_data
            else:
                db_session = SessionModel(id=session_id, data=session_data)
                db.add(db_session)
            db.commit()
        finally:
            db.close()
        
        # Set session cookie if it's a new session
        if getattr(request.state, "new_session", False):
            response.set_cookie(
                key="session_id",
                value=session_id,
                httponly=True,
                max_age=3600,
                samesite="lax"
            )
        
        return response

app.add_middleware(CustomSessionMiddleware)



# with open("config.json", "r") as f:
#     config = json.load(f)

# api_key = config["OPENROUTER_API_KEY"]
api_key = os.getenv('OPENROUTER_API_KEY')

# Set up templates directory
templates_directory = Path(__file__).parent / "templates"
templates = Jinja2Templates(directory=str(templates_directory))

# Create a temporary folder for file uploads
UPLOAD_FOLDER = "/tmp/resume_uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None


def enhance_resume(resume_text):
    try:
        system_prompt = {
            "role" : "system",
            "content": """You are a professional resume enhancement expert. 
            Your task is to improve the provided resume text while COMPLETELY PRESERVING ALL of the original content and sections.

            What You Should Do:
            Improve readability, conciseness, and clarity without changing the meaning.
            Make the writing more polished and professional while keeping it natural.
            
            What You Should NOT Do:
            Don't add anything new—not even for credibility.
            Don't change the order, structure, or formatting in any way.
            Don't remove anything, even if it seems unnecessary.
            Don't rewrite sentences in a way that alters the original intent.
            Your goal is to make the resume sound crisp, clear, and professional—without changing what's actually being said. Stick to the original content and just refine the language.

            Format your response with two clearly marked sections:
            <IMPROVED_RESUME>
            [The complete improved resume text with EVERY SINGLE element from the original preserved]
            </IMPROVED_RESUME>
            
            <CHANGES_MADE>
            [Bulleted list of specific language improvements made]
            </CHANGES_MADE>
            """
        }

        user_message = {
            "role" : "user",
            "content": "Here is my resume text to enhance. You MUST keep ALL sections, ALL projects, and ALL technical skills:\n\n{}".format(resume_text)
        }
        messages = [system_prompt, user_message]
        response = requests.post(
            url="https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": "Bearer {}".format(api_key),
                "Content-Type": "application/json",
            },
            
            data=json.dumps({
                "model": "deepseek/deepseek-r1:free",
                "messages": messages,
            })
        )
        # print("REsult:", response.json())
        result = response.json()["choices"][0]["message"]["content"]
        
        improved_resume_match = re.search(r"<IMPROVED_RESUME>\s*(.*?)\s*</IMPROVED_RESUME>", result, re.DOTALL)
        changes_made_match = re.search(r"<CHANGES_MADE>\s*(.*?)\s*</CHANGES_MADE>", result, re.DOTALL)

        if not improved_resume_match:
            print("⚠️ WARNING: <IMPROVED_RESUME> NOT FOUND IN RESPONSE!")

        if not changes_made_match:
            print("⚠️ WARNING: <CHANGES_MADE> NOT FOUND IN RESPONSE!")

        improved_resume_text = improved_resume_match.group(1) if improved_resume_match else "No improved resume found"
        changes_made_text = changes_made_match.group(1) if changes_made_match else "Removed Contact Information"
        
        # print("changes_made_text",changes_made_text)
        return {
            "improved_resume": improved_resume_text,
            "changes_made": changes_made_text,
            "error": False
        }
        
    except Exception as e:
        print(f"Exception:: {e}")
        return {
            "error": True,
            "message": f"Error: {str(e)}"
        }


def write_to_pdf(text):
    pdf_path = os.path.join(UPLOAD_FOLDER, "enhanced_resume.pdf")
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        leftMargin=50,
        rightMargin=50,
        topMargin=40,
        bottomMargin=40
    )
    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    centered_style = ParagraphStyle(
        'Centered',
        parent=styles['Normal'],
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        fontSize=14
    )
    
    bold_style = ParagraphStyle(
        'Bold',
        parent=styles['Normal'],
        fontName='Helvetica-Bold'
    )
    
    paragraphs = []
    
    for line in text.split("\n"):
        if not line.strip():
            paragraphs.append(Paragraph("<br/>", normal_style))
            continue
            
        line = re.sub(r'\*\*\*\*\*\*(.*?)\*\*\*\*\*\*', r'<b>\1</b>', line)
        
        line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
        
        if re.match(r'^<b>.*</b>$', line) and len(line) < 60: 
            clean_line = re.sub(r'</?b>', '', line)
            paragraphs.append(Paragraph(clean_line, centered_style))
        else:
            paragraphs.append(Paragraph(line, normal_style))
    
    doc.build(paragraphs)
    
    return pdf_path


def write_to_docx(text):

    docx_path = os.path.join(UPLOAD_FOLDER, "enhanced_resume.docx")
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)

    doc.save(docx_path)
    return docx_path



# Routes
@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index_flask.html", {"request": request})

@app.post("/upload")
async def upload_file(request: Request, resume: UploadFile = File(...), format: str = Form("pdf")):
    file_path = os.path.join(UPLOAD_FOLDER, resume.filename)
    
    # Save the uploaded file
    with open(file_path, "wb") as buffer:
        content = await resume.read()
        buffer.write(content)
    
    # Extract text based on file type
    if resume.filename.endswith('.pdf'):
        resume_text = extract_text_from_pdf(file_path)
        original_format = 'pdf'
    elif resume.filename.endswith('.docx'):
        resume_text = extract_text_from_docx(file_path)
        original_format = 'docx'
    else:
        return HTMLResponse("Unsupported file format. Please upload a PDF or DOCX file.")
    
    if not resume_text:
        return HTMLResponse("Failed to extract text from the uploaded file.")
    
    # Enhance the resume
    enhancement_result = enhance_resume(resume_text)
    
    # Store data in session
    improved_resume = enhancement_result.get('improved_resume')
    changes_made = enhancement_result.get('changes_made')

    request.state.session["improved_resume"] = improved_resume
    request.state.session["changes_made"] = changes_made


    request.state.session["output_format"] = format
    request.state.session["original_format"] = original_format

    # print("after storing",request.session)
    response = RedirectResponse(url="/results", status_code=303)
    return response


@app.get("/results", response_class=HTMLResponse)
async def show_results(request: Request):
    # print("In results",dict(request.state.session))
    if "improved_resume" not in request.state.session:
        # print("Improved resume not found")
        return RedirectResponse(url="/", status_code=303)
    
    changes_made = request.state.session.get("changes_made", "")
    if changes_made == "" or changes_made is None:
        return HTMLResponse(content="changes_made is not in session")
    formatted_changes = ""
    for line in changes_made.split('\n'):
        if line.strip():
            if line.strip().startswith('•') or line.strip().startswith('-') or line.strip().startswith('*'):
                formatted_changes += f"<li>{line.strip()[1:].strip()}</li>"
            else:
                formatted_changes += f"<li>{line.strip()}</li>"
    
    return templates.TemplateResponse(
        "results_flask.html", 
        {"request": request, "changes_made": formatted_changes}
    )


@app.get("/download")
async def download_file(request: Request):
    if "improved_resume" not in request.state.session:
        return RedirectResponse(url="/", status_code=303)
    
    improved_resume = request.state.session.get("improved_resume")
    output_format = request.state.session.get("output_format", "pdf")

    if output_format == 'docx':
        output_path = write_to_docx(improved_resume)
        filename = "enhanced_resume.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        output_path = write_to_pdf(improved_resume)
        filename = "enhanced_resume.pdf"
        media_type = "application/pdf"
    
    if not os.path.exists(output_path) or not os.access(output_path, os.R_OK):
        return HTMLResponse("Error: Could not generate the file. Please try again.")
    
    return FileResponse(
        path=output_path,
        filename=filename,
        media_type=media_type
    )


# Add a route to handle the placeholder image requests
@app.get("/api/placeholder/{width}/{height}")
async def placeholder_image(width: int, height: int):
    # Create a simple placeholder SVG
    svg_content = f"""
    <svg width="{width}" height="{height}" xmlns="http://www.w3.org/2000/svg">
        <rect width="100%" height="100%" fill="#e0e0e0" />
        <text x="50%" y="50%" font-family="Arial" font-size="24" fill="#666666" 
              text-anchor="middle" dominant-baseline="middle">
            {width} x {height}
        </text>
    </svg>
    """
    return Response(content=svg_content, media_type="image/svg+xml")